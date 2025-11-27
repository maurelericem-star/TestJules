# -*- coding: utf-8 -*-
import re
import json
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Fonctions de Parsing Spécifiques ---

def parse_header(doc, header_lines, styles):
    """
    Analyse les 3 premières lignes du Markdown (header) et les ajoute au document Word.
    """
    if len(header_lines) >= 3:
        p_name = doc.add_paragraph(header_lines[0].strip(), style=styles.get('nom_prenom', 'Normal'))
        p_title = doc.add_paragraph(header_lines[1].strip(), style=styles.get('titre_poste', 'Normal'))

        contact_line = header_lines[2].strip()
        p_contact = doc.add_paragraph(style=styles.get('coordonnees', 'Normal'))
        parse_rich_text(p_contact, contact_line)

        print("Header ajouté avec succès.")
    else:
        print("Avertissement: Le header Markdown ne contient pas les 3 lignes attendues.")

def parse_skills_table(doc, markdown_content, styles):
    """
    Analyse la section 'COMPÉTENCES CLÉS' du Markdown et remplit le tableau dans le document Word.
    """
    print("Début du remplissage du tableau de compétences...")
    try:
        skills_match = re.search(r'### COMPÉTENCES CLÉS\n([\s\S]*?)\n---', markdown_content)
        if not skills_match:
            print("Erreur: Section 'COMPÉTENCES CLÉS' introuvable dans le Markdown.")
            return

        skills_text = skills_match.group(1).strip()
        columns_data = re.split(r'\n#### ', skills_text)
        if columns_data:
            columns_data[0] = columns_data[0].replace("#### ","")

        parsed_columns = []
        for col_text in columns_data:
            lines = col_text.strip().split('\n')
            title = lines[0].strip()
            skills = [re.sub(r'^\s*\*\s*', '', line).strip() for line in lines[1:] if line.strip()]
            if title and skills:
                parsed_columns.append({'title': title, 'skills': skills})

        if not doc.tables:
            print("Erreur: Aucun tableau trouvé dans le template 'template_master.docx'.")
            return

        table = doc.tables[0]
        num_cols_to_fill = min(len(parsed_columns), len(table.columns))
        max_rows_needed = max(len(col['skills']) for col in parsed_columns) if parsed_columns else 0

        while len(table.rows) < max_rows_needed:
            table.add_row()

        for col_idx, col_data in enumerate(parsed_columns):
            if col_idx < num_cols_to_fill:
                for row_idx, skill in enumerate(col_data['skills']):
                    if row_idx < len(table.rows):
                        cell = table.cell(row_idx, col_idx)
                        cell.text = ''
                        p = cell.paragraphs[0]
                        p.text = skill
                        # Le style du tableau est défini sur la table elle-même, mais on peut le forcer ici.
                        # Le style mappé est celui de la table, pas du paragraphe.
                        # Laisser python-docx gérer le style hérité de la table est souvent mieux.

        print("Tableau de compétences rempli.")
    except Exception as e:
        print(f"Une erreur inattendue est survenue lors du traitement du tableau de compétences: {e}")


# --- Fonctions de Parsing de Contenu Générique ---

def add_hyperlink(paragraph, text, url):
    """
    Ajoute un hyperlien cliquable à un objet paragraphe.
    """
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    r = paragraph.add_run()
    r._r.append(hyperlink)
    r.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    r.font.underline = True
    return hyperlink

def parse_rich_text(paragraph, text):
    """
    Analyse une chaîne de caractères pour la syntaxe Markdown (gras, liens)
    et l'ajoute à un paragraphe existant avec le formatage correct.
    """
    pattern = r'\*\*(.*?)\*\*|\[(.*?)\]\((.*?)\)'

    last_end = 0
    for match in re.finditer(pattern, text):
        start, end = match.span()
        if start > last_end:
            paragraph.add_run(text[last_end:start])

        bold_text, link_text, link_url = match.groups()

        if bold_text:
            paragraph.add_run(bold_text).bold = True
        elif link_text and link_url:
            add_hyperlink(paragraph, link_text, link_url)

        last_end = end

    if last_end < len(text):
        paragraph.add_run(text[last_end:])

def parse_generic_section(doc, section_text, styles):
    """
    Analyse une section générique du CV (Profil, Expériences, etc.),
    ajoute son titre et son contenu formaté au document.
    """
    lines = section_text.strip().split('\n')
    if not lines:
        return

    title_line = lines[0]
    title_match = re.match(r'###\s*(.*)', title_line)

    if title_match:
        title = title_match.group(1).strip()
        doc.add_paragraph(title, style=styles.get('titre_section', 'Heading 3'))
        content_lines = lines[1:]
    else:
        content_lines = lines

    for line in content_lines:
        line = line.strip()
        if not line:
            continue

        style_name = styles.get('poste') if line.startswith('**') else 'Normal'
        if line.startswith('*'):
            style_name = 'List Bullet'
            line = re.sub(r'^\s*\*\s*', '', line).strip()

        p = doc.add_paragraph(style=style_name)
        parse_rich_text(p, line)

# --- Script Principal ---
def generate_cv():
    """
    Script principal pour générer le CV final à partir du template et du contenu Markdown.
    """
    try:
        # Étape 1: Charger le mapping des styles depuis le fichier JSON
        try:
            with open('styles.json', 'r', encoding='utf-8') as f:
                styles = json.load(f)
            print("Mapping des styles chargé depuis 'styles.json'.")
        except FileNotFoundError:
            print("Erreur: Le fichier 'styles.json' est introuvable. Veuillez d'abord exécuter 'analyzer_prep.py'.")
            return

        # Étape 2: Charger les documents de base
        doc = Document('template_master.docx')
        with open('MAUREL_Eric_CV_RATP_1_page.md', 'r', encoding='utf-8') as f:
            markdown_content = f.read()

        # Étape 3: Traiter le contenu
        header_lines = markdown_content.split('\n')[:3]
        parse_header(doc, header_lines, styles)

        parse_skills_table(doc, markdown_content, styles)

        main_content = '\n'.join(markdown_content.split('\n')[3:]).strip()
        sections = re.split(r'\n---\n', main_content)

        for section_text in sections:
            section_text = section_text.strip()
            if section_text.startswith('### COMPÉTENCES CLÉS'):
                continue
            if section_text:
                parse_generic_section(doc, section_text, styles)

        # Étape 4: Sauvegarder le document final
        output_filename = 'CV_Eric_Maurel_Genere.docx'
        doc.save(output_filename)
        print(f"CV généré avec succès sous le nom : {output_filename}")

    except FileNotFoundError as e:
        print(f"Erreur de fichier: {e}. Vérifiez que les fichiers de base sont présents.")
    except Exception as e:
        print(f"Une erreur est survenue: {e}")

if __name__ == '__main__':
    generate_cv()
