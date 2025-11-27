# -*- coding: utf-8 -*-
import logging
import json
from docx import Document

# --- Configuration du Logging ---
# Configure le logging pour écrire dans un fichier et sur la console.
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(message)s',
                    handlers=[logging.FileHandler("analysis_report.txt", mode='w'), logging.StreamHandler()])

def audit_manual_formatting(doc):
    """
    Scanne les paragraphes du document. Si un paragraphe utilise le style 'Normal'
    mais contient du formatage manuel (gras, italique, police forcée),
    il est loggué dans 'audit_styles.txt'.
    """
    logging.info("--- Début de l'Audit de Formatage Manuel ---")
    issues_found = False

    try:
        normal_style = doc.styles['Normal']
        style_font_name = normal_style.font.name
    except KeyError:
        logging.warning("Le style 'Normal' n'a pas été trouvé. L'audit sera moins précis.")
        style_font_name = None

    with open('audit_styles.txt', 'w', encoding='utf-8') as f:
        f.write("Rapport d'audit: Paragraphes avec formatage manuel sur style 'Normal':\n\n")

        for p in doc.paragraphs:
            if p.style and p.style.name == 'Normal':
                for run in p.runs:
                    # Vérifie le formatage direct (gras/italique) ou une police non standard
                    if run.bold or run.italic or (style_font_name and run.font.name and run.font.name != style_font_name):
                        issues_found = True
                        f.write(f"- \"{p.text}\"\n")
                        break # Inutile de vérifier les autres runs de ce paragraphe

    if issues_found:
        logging.info("Audit terminé : Problèmes de formatage trouvés. Voir 'audit_styles.txt'.")
    else:
        logging.info("Audit terminé : Aucun problème de formatage manuel trouvé.")
    logging.info("--- Fin de l'Audit ---")

def map_styles(doc):
    """
    Identifie et mappe les styles utilisés pour les éléments clés du CV.
    Retourne un dictionnaire de mapping pour une utilisation automatisée.
    """
    logging.info("\n--- Début du Mapping des Styles ---")
    styles_map = {}

    try:
        styles_map['nom_prenom'] = doc.paragraphs[0].style.name
        styles_map['titre_poste'] = doc.paragraphs[1].style.name
        styles_map['coordonnees'] = doc.paragraphs[2].style.name

        for p in doc.paragraphs:
            text_upper = p.text.upper()
            if 'PROFIL' in text_upper:
                styles_map['titre_section'] = p.style.name
            if 'BNP PARIBAS' in p.text:
                styles_map['poste'] = p.style.name

        if doc.tables:
            styles_map['tableau_competences'] = doc.tables[0].style.name
            # Fallback si le style est sur le paragraphe de la cellule
            if not styles_map['tableau_competences'] or styles_map['tableau_competences'] == 'Table Normal':
                 styles_map['tableau_competences'] = doc.tables[0].cell(0, 0).paragraphs[0].style.name

        for key, value in styles_map.items():
            logging.info(f"Style pour '{key}' -> '{value}'")
    except IndexError:
        logging.error("Le document semble vide ou mal structuré. Le mapping des styles a échoué.")

    logging.info("--- Fin du Mapping ---")
    return styles_map

def create_template(doc, skills_table_style):
    """
    Crée le fichier template en supprimant le contenu textuel non nécessaire
    tout en préservant les images, en-têtes, pieds de page et la structure
    du tableau de compétences.
    """
    logging.info("\n--- Début de la Création du Template Master ---")

    # 1. Identifier les tables à supprimer vs le tableau de compétences à vider
    tables_to_remove = []
    skills_table_found = False
    for table in doc.tables:
        # L'identification par le style de la table est plus robuste
        if table.style and table.style.name == skills_table_style:
            skills_table_found = True
            # Vider le contenu des cellules de ce tableau
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        p.clear()
        else:
            tables_to_remove.append(table)

    if skills_table_found:
        logging.info(f"Tableau de compétences (style '{skills_table_style}') identifié et vidé.")
    else:
        logging.warning(f"Le tableau de compétences avec le style '{skills_table_style}' n'a pas été trouvé.")

    # 2. Supprimer les tables indésirables
    for table in tables_to_remove:
        table._element.getparent().remove(table._element)

    # 3. Identifier les paragraphes à supprimer (ceux sans image)
    paragraphs_to_remove = []
    for p in doc.paragraphs:
        # La présence de 'graphicData' dans le XML du paragraphe indique une image
        if 'graphicData' not in p._p.xml:
            paragraphs_to_remove.append(p)

    # 4. Supprimer les paragraphes indésirables
    for p in paragraphs_to_remove:
        p._element.getparent().remove(p._element)

    logging.info("Contenu textuel supprimé, images et structure principale conservées.")
    doc.save('template_master.docx')
    logging.info("Template master sauvegardé sous : 'template_master.docx'")
    logging.info("--- Fin de la Création du Template ---")


# --- Script Principal ---
def main():
    try:
        # Utiliser la version .docx pour éviter les problèmes de compatibilité
        source_filename = 'MAUREL_Eric_CV_RATP.docx'
        doc = Document(source_filename)

        # Étape 1: Audit de la qualité des styles
        audit_manual_formatting(doc)

        # Étape 2: Mapping des styles pour l'automatisation
        style_mappings = map_styles(doc)

        # Étape 3: Sauvegarder le mapping pour le script generator.py
        with open('styles.json', 'w', encoding='utf-8') as f:
            json.dump(style_mappings, f, indent=4, ensure_ascii=False)
        logging.info("\nMapping des styles sauvegardé dans 'styles.json'")

        # Étape 4: Création du template master
        # Recharger le document pour s'assurer qu'il est dans son état original
        doc_for_template = Document(source_filename)
        skills_table_style = style_mappings.get('tableau_competences')

        if skills_table_style:
            create_template(doc_for_template, skills_table_style)
        else:
            logging.error("Impossible de créer le template car le style du tableau de compétences n'a pas été mappé.")

    except FileNotFoundError:
        logging.error(f"Erreur: Le fichier '{source_filename}' est introuvable.")
    except Exception as e:
        logging.error(f"Une erreur inattendue est survenue: {e}", exc_info=True)

if __name__ == '__main__':
    main()
